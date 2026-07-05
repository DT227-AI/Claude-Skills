#!/usr/bin/env python3
"""
build_report_docx.py — render an institutional equity investment due-diligence
report (content.json) to a formatted .docx, matching the fixed 第X部分 structure
used in bank/AIC 股权投资试点业务尽调报告.

Usage:
    python3 build_report_docx.py content.json output.docx

content.json schema
--------------------
{
  "meta": {
    "title_line1": "关于为北京XX科技有限公司办理",
    "title_line2": "不超过1亿元股权投资试点业务的尽调报告",
    "statement_paragraphs": [
      "我们已根据XX公司有关规章制度对纯股权项目进行调查，调查过程中做到了勤勉尽职。",
      "本尽职调查报告是我们在勤勉尽职的基础上所撰写，我们承诺报告不存在故意隐瞒负面消息、虚假记载，并对报告所涉及的数据和资料的真实性负责。"
    ],
    "signers": [
      {"role": "尽职调查人", "name": "张三、李四", "contact": "010-00000000"},
      {"role": "调查复核人", "name": "王五", "contact": "010-00000000"},
      {"role": "总经理", "name": "赵六", "contact": "010-00000000"},
      {"role": "经营主责任人", "name": "孙七", "contact": "025-00000000"}
    ],
    "submit_date": "2026年7月6日",
    "footer_confidential": "内部资料，请勿外传"
  },
  "parts": [
    {
      "label": "第一部分",
      "title": "股权投资项目方案",
      "blocks": [ <block>, ... ]
    },
    ...
  ]
}

Block types (each item in a "blocks" array is one of these):
  {"type": "h2", "text": "一、项目背景"}                     # 一/二/三 level heading
  {"type": "h3", "text": "（一）资产负债分析"}                 # (一)/(二) or subsection level
  {"type": "paragraph", "text": "..."}
  {"type": "bullets", "items": ["...", "..."]}
  {"type": "table", "headers": ["科目", "2021年", "2022年"], "rows": [["...", "...", "..."], ...]}
  {"type": "risk_item", "index": 1, "title": "...", "description": "...", "measure": "..."}
  {"type": "callout", "title": "...", "text": "..."}          # shaded emphasis box (e.g. 结论性意见 recommendation)

Do not invent new block types — the renderer only understands the six above.
"""
import json
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT_LIGHT = "1F3A5F"
CALLOUT_FILL = "EDF1F6"
HEADER_FILL = "1F3A5F"
FONT_LATIN = "Calibri"
FONT_CJK = "微软雅黑"
FONT_CJK_BODY = "宋体"


def set_cjk_font(run, latin=FONT_LATIN, cjk=FONT_CJK, size=None, bold=None, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def add_shading(cell_or_paragraph, hex_color, is_cell=True):
    el = cell_or_paragraph._tc if is_cell else cell_or_paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    if is_cell:
        tcPr = el.get_or_add_tcPr()
        tcPr.append(shd)
    else:
        el.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tblPr.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_footer(section, meta):
    text = meta.get("footer_confidential", "")
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    if text:
        r1 = p.add_run(text + "  ")
        set_cjk_font(r1, size=8, color=RGBColor(0x80, 0x80, 0x80))
    add_page_field(p)


def add_declaration_page(doc, meta):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("title_line1", "")), size=20, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("title_line2", "")), size=20, bold=True, color=ACCENT)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    set_cjk_font(p.add_run("声明与保证："), size=11.5, bold=True)
    doc.add_paragraph()

    for para_text in meta.get("statement_paragraphs", []):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(8)
        set_cjk_font(p.add_run(para_text), size=10.5, cjk=FONT_CJK_BODY)

    doc.add_paragraph()
    for signer in meta.get("signers", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        line = f"{signer.get('role', '')}： {signer.get('name', '')}"
        if signer.get("contact"):
            line += f" 联系方式：{signer['contact']}"
        set_cjk_font(p.add_run(line), size=10.5, cjk=FONT_CJK_BODY)

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_cjk_font(p.add_run(f"提交日期： {meta.get('submit_date', '')}"), size=10.5, cjk=FONT_CJK_BODY)

    doc.add_page_break()


def add_part_title(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    set_cjk_font(p.add_run(f"{label} {title}"), size=17, bold=True, color=ACCENT)
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), ACCENT_LIGHT)
    border.append(bottom)
    pPr.append(border)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    set_cjk_font(p.add_run(text), size=13.5, bold=True, color=ACCENT)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    set_cjk_font(p.add_run(text), size=11.5, bold=True)


def add_paragraph_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    set_cjk_font(p.add_run(text), size=10.5, cjk=FONT_CJK_BODY)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        set_cjk_font(p.add_run(item), size=10.5, cjk=FONT_CJK_BODY)


def add_table_block(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cjk_font(p.add_run(str(h)), size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        add_shading(hdr_cells[i], HEADER_FILL, is_cell=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_cjk_font(p.add_run(str(val)), size=9.5, cjk=FONT_CJK_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_risk_item(doc, index, title, description, measure):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_cjk_font(p.add_run(f"{index}. {title}"), size=11, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    set_cjk_font(p.add_run(description), size=10.5, cjk=FONT_CJK_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    set_cjk_font(p.add_run("应对措施："), size=10.5, bold=True, cjk=FONT_CJK_BODY)
    set_cjk_font(p.add_run(measure), size=10.5, cjk=FONT_CJK_BODY)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    add_shading(cell, CALLOUT_FILL, is_cell=True)
    p0 = cell.paragraphs[0]
    set_cjk_font(p0.add_run(title), size=10.5, bold=True, color=ACCENT)
    p1 = cell.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    set_cjk_font(p1.add_run(text), size=10.5, cjk=FONT_CJK_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def render_blocks(doc, blocks):
    for b in blocks:
        t = b.get("type")
        if t == "h2":
            add_h2(doc, b["text"])
        elif t == "h3":
            add_h3(doc, b["text"])
        elif t == "paragraph":
            add_paragraph_block(doc, b["text"])
        elif t == "bullets":
            add_bullets(doc, b["items"])
        elif t == "table":
            add_table_block(doc, b["headers"], b["rows"])
        elif t == "risk_item":
            add_risk_item(doc, b["index"], b["title"], b["description"], b["measure"])
        elif t == "callout":
            add_callout(doc, b["title"], b["text"])
        else:
            raise ValueError(
                f"Unknown block type: {t!r}. Allowed: h2, h3, paragraph, bullets, table, risk_item, callout."
            )


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 build_report_docx.py content.json output.docx")
        sys.exit(1)
    content_path, output_path = sys.argv[1], sys.argv[2]
    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    meta = content["meta"]
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_LATIN
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK_BODY)

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        build_footer(section, meta)

    add_declaration_page(doc, meta)

    for part in content["parts"]:
        add_part_title(doc, part["label"], part["title"])
        render_blocks(doc, part["blocks"])
        doc.add_page_break()

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
