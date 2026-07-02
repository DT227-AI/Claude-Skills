#!/usr/bin/env python3
"""
build_report_docx.py — render an investment deep-dive profile (content.json) to a formatted .docx.

Usage:
    python3 build_report_docx.py content.json output.docx

content.json schema
--------------------
{
  "meta": {
    "org_label": "工银投资 · 国创引导基金专项柔性团队",   # top-of-cover small label
    "company_name": "SimpleX",                              # big cover title
    "report_type": "深度分析画像",
    "tagline": "AI for Science · Bio LLM 基座公司",         # one-line domain tag on cover
    "subtitle": "深度画像",
    "date_label": "2026年6月（6/29更新）",
    "confidentiality_label": "工银投资 · 内部资料",
    "disclaimer_cover": "本材料仅供工银投资内部参阅，不构成投资建议",
    "footer_confidential": "本材料仅供工银投资内部参阅，不得外传",
    "footer_disclaimer": "本材料仅供工银投资内部参阅，不构成投资建议"
  },
  "sections": [
    {
      "number": "一",                 # 一/二/三/四/五 ...
      "title": "公司概览与定位",
      "blocks": [ <block>, ... ]
    },
    ...
  ],
  "appendix": {
    "title": "附录：关键数据速查",
    "blocks": [ <block>, ... ]
  }
}

Block types (each item in a "blocks" array is one of these):
  {"type": "subheading", "text": "1.1 一句话定位"}
  {"type": "paragraph", "text": "..."}
  {"type": "bullets", "items": ["...", "..."]}
  {"type": "table", "headers": ["维度", "信息"], "rows": [["公司全称", "..."], ...]}
  {"type": "callout", "title": "团队判断", "text": "..."}   # shaded synthesis box

Do not invent new block types — the renderer only understands the five above.
"""
import json
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)      # dark navy, headings
ACCENT_LIGHT = "1F3A5F"                  # hex for shading
CALLOUT_FILL = "EDF1F6"                  # light blue-gray for synthesis boxes
HEADER_FILL = "1F3A5F"                   # table header row fill
FONT_LATIN = "Calibri"
FONT_CJK = "微软雅黑"
FONT_CJK_BODY = "宋体"


def set_cjk_font(run, latin=FONT_LATIN, cjk=FONT_CJK, size=None, bold=None, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def add_shading(cell_or_paragraph, hex_color, is_cell=True):
    el = cell_or_paragraph._tc if is_cell else cell_or_paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    if is_cell:
        tcPr = el.get_or_add_tcPr()
        tcPr.append(shd)
    else:
        el.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tblPr.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "目录生成中，请在 Word 中按 Ctrl+A 后 F9 更新域，或右键选择「更新域」。"
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    t_run = OxmlElement('w:r')
    t_run.append(placeholder)
    run._r.addnext(t_run)
    run._r.append(fld_end)


def build_footer(section, meta):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r1 = p.add_run(meta.get("footer_confidential", ""))
    set_cjk_font(r1, size=8, color=RGBColor(0x80, 0x80, 0x80))
    p.add_run("\n")
    add_page_field(p)
    r2 = p.add_run(f"  {meta.get('footer_disclaimer', '')}")
    set_cjk_font(r2, size=8, color=RGBColor(0x80, 0x80, 0x80))


def add_cover(doc, meta):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("org_label", "")), size=12, bold=True, color=ACCENT)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("company_name", "")), size=40, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("report_type", "")), size=20, bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("tagline", "")), size=13, color=RGBColor(0x40, 0x40, 0x40))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("subtitle", "")), size=13, color=RGBColor(0x40, 0x40, 0x40))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("date_label", "")), size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("confidentiality_label", "")), size=11, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(p.add_run(meta.get("disclaimer_cover", "")), size=9, color=RGBColor(0x80, 0x80, 0x80))

    doc.add_page_break()


def add_toc(doc):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cjk_font(h.add_run("目 录"), size=22, bold=True, color=ACCENT)
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_toc_field(p)
    doc.add_page_break()


def add_section_title(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.border_bottom = None
    set_cjk_font(p.add_run(f"{number}、{title}"), size=18, bold=True, color=ACCENT)
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), ACCENT_LIGHT)
    border.append(bottom)
    pPr.append(border)


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_cjk_font(p.add_run(text), size=13, bold=True, color=ACCENT)


def add_paragraph_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    set_cjk_font(p.add_run(text), size=10.5, cjk=FONT_CJK_BODY)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        set_cjk_font(p.add_run(item), size=10.5, cjk=FONT_CJK_BODY)


def add_table_block(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        set_cjk_font(p.add_run(h), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        add_shading(hdr_cells[i], HEADER_FILL, is_cell=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_cjk_font(p.add_run(str(val)), size=10, cjk=FONT_CJK_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    add_shading(cell, CALLOUT_FILL, is_cell=True)
    p0 = cell.paragraphs[0]
    set_cjk_font(p0.add_run(title), size=10.5, bold=True, color=ACCENT)
    p1 = cell.add_paragraph()
    p1.paragraph_format.line_spacing = 1.25
    set_cjk_font(p1.add_run(text), size=10.5, cjk=FONT_CJK_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def render_blocks(doc, blocks):
    for b in blocks:
        t = b.get("type")
        if t == "subheading":
            add_subheading(doc, b["text"])
        elif t == "paragraph":
            add_paragraph_block(doc, b["text"])
        elif t == "bullets":
            add_bullets(doc, b["items"])
        elif t == "table":
            add_table_block(doc, b["headers"], b["rows"])
        elif t == "callout":
            add_callout(doc, b["title"], b["text"])
        else:
            raise ValueError(f"Unknown block type: {t!r}. Allowed: subheading, paragraph, bullets, table, callout.")


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 build_report_docx.py content.json output.docx")
        sys.exit(1)
    content_path, output_path = sys.argv[1], sys.argv[2]
    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    meta = content["meta"]
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_LATIN
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK_BODY)

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        build_footer(section, meta)

    add_cover(doc, meta)
    add_toc(doc)

    for sec in content["sections"]:
        add_section_title(doc, sec["number"], sec["title"])
        render_blocks(doc, sec["blocks"])
        doc.add_page_break()

    appendix = content.get("appendix")
    if appendix:
        p = doc.add_paragraph()
        set_cjk_font(p.add_run(appendix["title"]), size=18, bold=True, color=ACCENT)
        render_blocks(doc, appendix["blocks"])

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
